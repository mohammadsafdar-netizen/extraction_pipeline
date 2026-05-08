"""
Unified extraction runner for the Input/ folder.

Handles the four document classes you typically get with an insurance
submission package:

  - .pdf  ACORD application, supplemental, loss run, community map →
          existing merged or loss-run pipeline (or pure-VLM for unknowns)
  - .docx narrative attachments / emails / underwriter notes →
          python-docx structured text extraction
  - .xls  / .xlsx Statement of Values, exposure schedules →
          pandas-based per-sheet/row extraction

Output: one JSON file per input doc in input_extracted/, plus a master
input_extracted/ALL.json with everything in a single file.
"""
import argparse
import io
import json
import re
import shutil
import subprocess
import sys
from pathlib import Path

import fitz  # for PDF page count + classification
import pandas as pd

# Lazy-import the heavy pipelines only when needed
REPO = Path(__file__).resolve().parent
INPUT_DIR_DEFAULT = REPO / "input_docs" / "Input"
OUT_DIR_DEFAULT = REPO / "input_extracted"
PDF_DIR = REPO / "pdfs"


# ── docx ──

def extract_docx(path: Path) -> dict:
    """Extract structured content from a .docx file:
       - paragraphs (with style)
       - tables (rows of cells)
       - section count
       - inline images (count only)"""
    from docx import Document

    doc = Document(str(path))
    paragraphs = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        paragraphs.append({
            "style": p.style.name if p.style else "Normal",
            "text": text,
        })

    tables = []
    for ti, t in enumerate(doc.tables):
        rows = []
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                rows.append(cells)
        if rows:
            tables.append({"index": ti, "rows": rows})

    n_images = sum(1 for s in doc.inline_shapes)

    return {
        "source_file": path.name,
        "type": "docx",
        "n_paragraphs": len(paragraphs),
        "n_tables": len(tables),
        "n_inline_images": n_images,
        "paragraphs": paragraphs,
        "tables": tables,
    }


# ── xls / xlsx ──

def extract_excel(path: Path) -> dict:
    """Extract every sheet of an Excel file as a list of row records.
       Uses pandas with openpyxl (xlsx) or xlrd (xls)."""
    suffix = path.suffix.lower()
    engine = "openpyxl" if suffix == ".xlsx" else "xlrd"

    out_sheets = []
    try:
        xls = pd.ExcelFile(str(path), engine=engine)
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name, header=None,
                                dtype=object)
            # Drop fully-empty rows + columns
            df = df.dropna(how="all", axis=0).dropna(how="all", axis=1)
            df = df.fillna("")

            # Find first non-empty row to use as header
            rows = df.values.tolist()
            header = None
            data_rows = []
            for r in rows:
                if header is None and any(str(c).strip() for c in r):
                    header = [str(c).strip() for c in r]
                    continue
                if header:
                    record = {}
                    for k, v in zip(header, r):
                        if k:
                            record[k] = ("" if v == "" else
                                         (v.isoformat() if hasattr(v, "isoformat")
                                          else str(v).strip()))
                    data_rows.append(record)

            out_sheets.append({
                "sheet_name": sheet_name,
                "n_rows": len(rows),
                "header": header or [],
                "records": data_rows,
                "raw_grid": rows,  # full grid (untruncated)
            })
    except Exception as e:
        return {
            "source_file": path.name,
            "type": "excel",
            "_error": f"{type(e).__name__}: {e}",
        }

    return {
        "source_file": path.name,
        "type": "excel",
        "n_sheets": len(out_sheets),
        "sheets": out_sheets,
    }


# ── pdf classification + dispatch ──

LOSS_RUN_PATTERNS = re.compile(r"\b(loss[\s_-]?run|LR[_\s-]|claims?[\s_-]history)",
                                re.IGNORECASE)


_ACORD_FORM_NUMBERS = re.compile(
    r"\bacord[_\s-]*(125|126|127|130|131|137|140|163|823)\b", re.IGNORECASE)
_ACORD_APP_HINT = re.compile(r"\b(acord.+app|app.+acord|application)\b",
                              re.IGNORECASE)


def classify_pdf(path: Path) -> str:
    """Filename-based classification (legacy fast-path).
       Returns one of {acord_application, loss_run, supplemental, other}.

       For content-based per-page classification — including mixed PDFs
       with both ACORD pages AND loss-run pages — use
       detect_form_type.classify_pdf_pages() instead."""
    name = path.name.lower()
    if LOSS_RUN_PATTERNS.search(name) or name.startswith("lr_"):
        return "loss_run"
    if _ACORD_FORM_NUMBERS.search(name) or _ACORD_APP_HINT.search(name):
        return "acord_application"
    if "supp" in name or "supplemental" in name:
        return "supplemental"
    if "questionnaire" in name:
        return "questionnaire"
    return "other"


def classify_pdf_by_content(path: Path, templates_dir: Path) -> dict:
    """Content-based classification. Returns:
         {primary_kind, n_pages, kinds_present: {ACORD_FORM, LOSS_RUN, ...},
          page_groups: [{kind, pages}, ...]}
       Uses detect_form_type.classify_pdf_pages so multi-form PDFs get
       routed page-by-page to the right pipeline."""
    from detect_form_type import classify_pdf_pages
    result = classify_pdf_pages(path, templates_dir)
    kinds = {}
    for entry in result["pages"].values():
        kinds[entry["kind"]] = kinds.get(entry["kind"], 0) + 1
    # Primary kind = most-frequent non-EMPTY/non-UNKNOWN kind
    informative = {k: v for k, v in kinds.items()
                   if k not in ("EMPTY", "UNKNOWN")}
    primary = max(informative, key=informative.get) if informative else "UNKNOWN"
    return {
        "primary_kind": primary,
        "n_pages": result["n_pages"],
        "kinds_present": kinds,
        "page_groups": result["groups"],
    }


def extract_pdf(path: Path, out_dir: Path) -> dict:
    """Route PDF to appropriate pipeline:
       - acord_application → merged bbox+VLM (run_vllm_merged.process_pdf)
                             with auto-detected page_map
       - loss_run → pdfplumber+VLM cross-check (run_vllm_loss_runs.process_pdf)
       - other → pure-VLM extraction
       For mixed PDFs (e.g. ACORD app pages + embedded loss-run pages),
       falls through to acord_application using the per-page detection
       which will skip non-template pages."""
    doc_type = classify_pdf(path)
    print(f"  type (filename): {doc_type}")

    # Content-based classification — overrides filename if signal is strong
    try:
        content = classify_pdf_by_content(path, REPO / "templates")
        primary = content["primary_kind"]
        kinds_summary = ", ".join(f"{k}={v}" for k, v in content["kinds_present"].items())
        print(f"  type (content): {primary}  ({kinds_summary})")
        if primary == "ACORD_FORM" and doc_type != "acord_application":
            doc_type = "acord_application"
        elif primary == "LOSS_RUN" and doc_type != "loss_run":
            doc_type = "loss_run"
    except Exception as e:
        print(f"  (content-classify skipped: {e})")

    # Ensure the PDF is in pdfs/ since the pipeline runners read from there
    pdf_target = PDF_DIR / path.name
    if not pdf_target.exists():
        PDF_DIR.mkdir(exist_ok=True)
        shutil.copy(str(path), str(pdf_target))

    if doc_type == "acord_application":
        # Always auto-detect form type per page when called from the input
        # extractor — handles arbitrary multi-form ACORD packets correctly.
        # The hardcoded PDFS dict in run_vllm_merged.py is for the original
        # test-set regression fixture only; here we prefer content-based
        # detection so submissions with the same filename but different
        # contents (or with more ACORD pages than the legacy hardcoded
        # 4-page map covers) get extracted correctly.
        from run_vllm_merged import (process_pdf as merged_process,
                                       PAGE_MAP_125_ONLY)
        from detect_form_type import detect_form_type_for_pdf
        page_map = detect_form_type_for_pdf(pdf_target, REPO / "templates")
        if page_map:
            kinds = sorted({tmpl for tmpl, _ in page_map.values()})
            print(f"  auto-detected {len(page_map)} ACORD pages: {kinds}")
        else:
            # Fall back to assuming first 4 pages are ACORD 125
            page_map = PAGE_MAP_125_ONLY
            print(f"  no template matches; falling back to ACORD-125 1-4")
        return merged_process(
            path.name, "acord_application", page_map,
            PDF_DIR, REPO / "templates", dpi=150,
        )

    if doc_type == "loss_run":
        from run_vllm_loss_runs import process_pdf as loss_process
        return loss_process(path.name, PDF_DIR, dpi=150)

    # Other: render each page, send to VLM with GENERALIZED prompt
    print(f"  using pure-VLM extraction for unknown PDF type")
    from run_vllm_merged import run_vlm_extraction
    vlm_result, total_pages = run_vlm_extraction(
        path, doc_type or "other", dpi=150,
        page_map=None, templates_dir=REPO / "templates",
    )
    return {
        "source_file": path.name,
        "type": "pdf",
        "document_type": doc_type,
        "total_pages": total_pages,
        "model": "qwen3-vl-8b",
        "pages": [
            {"page": pg, "data": vlm_result.get(pg)}
            for pg in sorted(vlm_result)
        ],
    }


# ── main ──

def safe_name(name: str) -> str:
    return re.sub(r"[^A-Za-z0-9._-]+", "_", name).strip("_")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--input-dir", default=str(INPUT_DIR_DEFAULT))
    ap.add_argument("--out", default=str(OUT_DIR_DEFAULT))
    ap.add_argument("--overwrite", action="store_true")
    args = ap.parse_args()

    input_dir = Path(args.input_dir)
    out_dir = Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)

    files = sorted([f for f in input_dir.rglob("*") if f.is_file()])
    if not files:
        print(f"No files in {input_dir}")
        return 1

    all_extractions = []
    for path in files:
        suffix = path.suffix.lower()
        out_path = out_dir / f"{safe_name(path.stem)}.json"

        print(f"\n{'='*70}\n{path.name}  ({suffix})")
        if out_path.exists() and not args.overwrite:
            print(f"  SKIP: {out_path.name} already exists")
            with open(out_path) as f:
                all_extractions.append(json.load(f))
            continue

        try:
            if suffix == ".pdf":
                result = extract_pdf(path, out_dir)
            elif suffix == ".docx":
                result = extract_docx(path)
            elif suffix in (".xls", ".xlsx"):
                result = extract_excel(path)
            else:
                result = {"source_file": path.name, "type": suffix.lstrip("."),
                           "_error": "unsupported file type"}
        except Exception as e:
            import traceback
            traceback.print_exc()
            result = {"source_file": path.name, "type": suffix.lstrip("."),
                       "_error": f"{type(e).__name__}: {e}"}

        with open(out_path, "w") as f:
            json.dump(result, f, indent=2, default=str)
        print(f"  -> {out_path.name}")
        all_extractions.append(result)

    master = {
        "input_dir": str(input_dir),
        "n_files": len(all_extractions),
        "by_type": {},
        "files": all_extractions,
    }
    for r in all_extractions:
        t = r.get("type") or r.get("document_type") or "unknown"
        master["by_type"][t] = master["by_type"].get(t, 0) + 1

    with open(out_dir / "ALL.json", "w") as f:
        json.dump(master, f, indent=2, default=str)
    print(f"\n{'='*70}\nSUMMARY")
    print(f"  total: {len(all_extractions)} files")
    for t, n in master["by_type"].items():
        print(f"    {t}: {n}")
    print(f"  master: {out_dir / 'ALL.json'}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
